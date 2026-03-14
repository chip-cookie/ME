import io
import re
import logging
import zipfile
import subprocess
import chardet
from pathlib import Path
from typing import Callable

from docling.document_converter import DocumentConverter

log = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# 포맷별 파서 — 각각 (file_path: str) → str 반환, 실패 시 ""
# ---------------------------------------------------------------------------

def _parse_hwp_cli(file_path: str) -> str:
    """hwp5txt CLI (pyhwp 패키지 설치 시 제공)"""
    try:
        r = subprocess.run(
            ["hwp5txt", file_path],
            capture_output=True, text=True, timeout=60,
        )
        return r.stdout.strip() if r.returncode == 0 else ""
    except Exception:
        return ""


def _parse_hwp_api(file_path: str) -> str:
    """pyhwp Python API를 직접 사용하여 텍스트 추출"""
    try:
        from hwp5.xmlmodel import Hwp5File
        buf = io.StringIO()
        with Hwp5File(file_path) as hwp:
            for paragraph in hwp.bodytext.paragraphs():
                text = paragraph.get_text()
                if text:
                    buf.write(text + "\n")
        return buf.getvalue().strip()
    except Exception:
        return ""


def _parse_hwpx(file_path: str) -> str:
    """
    .hwpx = ZIP 기반 XML 포맷 (한글 2014 이후).
    Contents/section*.xml 파일에서 텍스트를 추출합니다.
    """
    try:
        with zipfile.ZipFile(file_path, "r") as zf:
            names = zf.namelist()
            section_files = sorted(
                [n for n in names if re.match(r"Contents/[Ss]ection\d+\.xml", n)]
            )
            if not section_files:
                section_files = sorted(
                    [n for n in names if n.endswith(".xml") and "section" in n.lower()]
                )

            text_parts = []
            for sec in section_files:
                xml_str = zf.read(sec).decode("utf-8", errors="replace")
                texts = re.findall(r"<(?:hp:)?t[^>]*>(.*?)</(?:hp:)?t>", xml_str, re.DOTALL)
                cleaned = " ".join(t.strip() for t in texts if t.strip())
                if cleaned:
                    text_parts.append(cleaned)

        return "\n\n".join(text_parts).strip()
    except Exception:
        return ""


def _parse_pdf_pdfplumber(file_path: str) -> str:
    """pdfplumber — 표·컬럼 레이아웃에 강함 (Docling 실패 시 1차 폴백)"""
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text(x_tolerance=2, y_tolerance=2) or ""
                if text.strip():
                    parts.append(text.strip())
        return "\n\n".join(parts)
    except Exception:
        return ""


def _parse_pdf_pypdf(file_path: str) -> str:
    """pypdf — 단순 텍스트 추출 (최후 수단)"""
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        parts = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(p.strip() for p in parts if p.strip())
    except Exception:
        return ""


def _parse_docx_python_docx(file_path: str) -> str:
    """
    python-docx — 단락·표·스타일을 Markdown으로 변환.
    Docling이 실패했을 때 폴백.
    """
    try:
        from docx import Document
        from docx.table import Table as DocxTable

        doc   = Document(file_path)
        parts = []

        for block in doc.element.body:
            tag = block.tag.split("}")[-1]

            if tag == "p":
                para_obj = next((p for p in doc.paragraphs if p._element is block), None)
                if not para_obj:
                    continue
                text  = para_obj.text.strip()
                if not text:
                    continue
                style = para_obj.style.name if para_obj.style else ""
                if   "Heading 1" in style: parts.append(f"# {text}")
                elif "Heading 2" in style: parts.append(f"## {text}")
                elif "Heading 3" in style: parts.append(f"### {text}")
                else:                      parts.append(text)

            elif tag == "tbl":
                tbl = DocxTable(block, doc)
                if not tbl.rows:
                    continue
                rows_md = []
                for i, row in enumerate(tbl.rows):
                    cells = [c.text.replace("|", "｜").strip() for c in row.cells]
                    rows_md.append("| " + " | ".join(cells) + " |")
                    if i == 0:
                        rows_md.append("|" + "|".join(["---"] * len(cells)) + "|")
                parts.append("\n".join(rows_md))

        return "\n\n".join(parts)
    except Exception:
        return ""


def _parse_plain_text(file_path: str) -> str:
    """
    .txt / .md — chardet으로 인코딩 감지.
    한국어 파일은 EUC-KR / CP949 / UTF-8 혼용 가능.
    """
    try:
        raw      = Path(file_path).read_bytes()
        detected = chardet.detect(raw)
        enc      = (detected.get("encoding") or "utf-8").upper()
        if enc in ("EUC-KR", "CP949", "EUC_KR", "JOHAB"):
            enc = "cp949"
        else:
            enc = enc.lower()
        return raw.decode(enc, errors="replace").strip()
    except Exception:
        return ""


def _parse_with_docling(file_path: str) -> str:
    """Docling — PDF/DOCX/DOC/ODT/RTF 등 범용 파서"""
    try:
        result = DocumentConverter().convert(file_path)
        return result.document.export_to_markdown()
    except Exception:
        return ""


# ---------------------------------------------------------------------------
# 형식별 파서 체인 (순서대로 시도, 첫 번째 비어있지 않은 결과 사용)
# ---------------------------------------------------------------------------

_PARSER_CHAINS: dict[str, list[Callable[[str], str]]] = {
    # 한글
    ".hwp":  [_parse_hwp_cli, _parse_hwp_api, _parse_with_docling],
    ".hwpx": [_parse_hwpx, _parse_with_docling],
    # PDF
    ".pdf":  [_parse_with_docling, _parse_pdf_pdfplumber, _parse_pdf_pypdf],
    # Word
    ".docx": [_parse_with_docling, _parse_docx_python_docx],
    ".doc":  [_parse_with_docling],
    # 기타 오피스
    ".odt":  [_parse_with_docling],
    ".rtf":  [_parse_with_docling],
    ".pptx": [_parse_with_docling],
    ".xlsx": [_parse_with_docling],
    # 텍스트
    ".txt":  [_parse_plain_text, _parse_with_docling],
    ".md":   [_parse_plain_text],
    ".text": [_parse_plain_text],
}

SUPPORTED_EXTENSIONS = set(_PARSER_CHAINS.keys())


# ---------------------------------------------------------------------------
# 메인 서비스
# ---------------------------------------------------------------------------

class ParserService:

    async def parse_document(self, file_path: str) -> dict:
        ext  = Path(file_path).suffix.lower()
        text = self._dispatch(file_path, ext)

        if not text:
            log.warning("모든 파서 실패: %s", file_path)

        structured = self._structure_markdown(text)
        return {
            "raw_markdown": text,
            "sections": structured["sections"],
            "ext": ext,
            "char_count": len(text),
        }

    # ------------------------------------------------------------------

    def _dispatch(self, file_path: str, ext: str) -> str:
        chain = _PARSER_CHAINS.get(ext)
        if chain is None:
            log.warning("지원하지 않는 확장자 '%s' — Docling+plain 텍스트로 시도", ext)
            chain = [_parse_with_docling, _parse_plain_text]

        for parser in chain:
            name = parser.__name__
            try:
                result = parser(file_path)
                if result and result.strip():
                    log.info(
                        "파싱 성공 [%s]: %s (%d자)",
                        name, Path(file_path).name, len(result),
                    )
                    return result
                log.debug("빈 결과 [%s], 다음 파서 시도", name)
            except Exception as exc:
                log.warning("[%s] 예외: %s", name, exc)

        return ""

    # ------------------------------------------------------------------

    def _structure_markdown(self, markdown_text: str) -> dict:
        """한국어 자기소개서 섹션 헤더를 인식하여 청크로 분리"""
        if not markdown_text.strip():
            return {"sections": []}

        pattern = re.compile(
            r"(?m)^(?:\s*#{1,3}\s*)?(?:"
            r"\[?\d+\]?[\.\)]\s+.+"
            r"|성장\s*과정.+"
            r"|지원\s*동기.+"
            r"|입사\s*후\s*포부.+"
            r"|자기\s*소개.+"
            r"|직무\s*역량.+"
            r"|경력\s*사항.+"
            r"|학력\s*사항.+"
            r"|자격\s*증.+"
            r")$"
        )
        parts   = pattern.split(markdown_text)
        headers = pattern.findall(markdown_text)

        sections = []
        body_parts = parts[1:] if len(parts) > 1 else parts
        for i, body in enumerate(body_parts):
            title = headers[i].strip() if i < len(headers) else f"section_{i+1}"
            for idx, chunk in enumerate(self._chunk_text_with_overlap(body.strip())):
                if chunk.strip():
                    sections.append({"title": f"{title}_part{idx+1}", "content": chunk})

        if not sections and markdown_text.strip():
            for idx, chunk in enumerate(self._chunk_text_with_overlap(markdown_text.strip())):
                sections.append({"title": f"section_part{idx+1}", "content": chunk})

        return {"sections": sections}

    def _chunk_text_with_overlap(
        self, text: str, max_length: int = 800, overlap: int = 150
    ) -> list[str]:
        if len(text) <= max_length:
            return [text]
        chunks, start = [], 0
        while start < len(text):
            end = start + max_length
            if end >= len(text):
                chunks.append(text[start:])
                break
            safe_end = text.rfind(". ", start + 400, end)
            if safe_end == -1:
                safe_end = text.rfind(" ", start + 600, end)
            actual_end = (safe_end + 1) if safe_end != -1 else end
            chunks.append(text[start:actual_end])
            start = actual_end - overlap
        return chunks
