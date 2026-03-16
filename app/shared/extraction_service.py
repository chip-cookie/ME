"""
Dual-Path Extraction Service

This module provides high-accuracy text extraction from documents using
an ensemble approach:
1. Layer 1 (Docling): Structural extraction with layout understanding
2. Layer 2 (Native): Fast extraction using specialized libraries
3. Similarity-based reconciliation with optional LLM repair
"""

import os
import logging
from typing import Optional, Tuple, Dict, Any, List
from pathlib import Path

# Layer 1: Docling
# Temporarily disabled due to torch DLL load error (WinError 1114)
DOCLING_AVAILABLE = False

# Layer 2: Native Parsers
try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# HWP (한글) Support
try:
    from pyhwpx import Hwp
    HWP_AVAILABLE = True
except ImportError:
    try:
        from hwp_extract import Hwp5File
        HWP_AVAILABLE = True
    except ImportError:
        HWP_AVAILABLE = False

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# Similarity Calculation
from rapidfuzz import fuzz

# LLM for repair
try:
    from langchain_community.llms import Ollama
    OLLAMA_AVAILABLE = True
except ImportError:
    OLLAMA_AVAILABLE = False

from app.core.config import get_settings

settings = get_settings()
logger = logging.getLogger(__name__)


class ExtractionResult:
    """Container for extraction results including text, images, and metadata."""
    def __init__(
        self,
        text: str = "",
        images: List[Dict[str, Any]] = None,
        metadata: Dict[str, Any] = None,
        source: str = "unknown"
    ):
        self.text = text
        self.images = images or []
        self.metadata = metadata or {}
        self.source = source  # 'docling', 'native', 'merged'
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "text": self.text,
            "images": self.images,
            "metadata": self.metadata,
            "source": self.source
        }


class DualPathExtractor:
    """
    Dual-Path Extraction Pipeline for high-accuracy document processing.
    
    Strategy:
    1. Extract using Docling (layout-aware, handles tables/images)
    2. Extract using native libraries (fast, robust for raw text)
    3. Compare similarity and merge/repair if needed
    """
    
    SIMILARITY_THRESHOLD = 0.85  # Below this, trigger LLM repair
    IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.tiff'}
    DOCUMENT_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.html', '.pptx'}
    
    def __init__(self, llm: Optional[Any] = None):
        """Initialize extractor with optional LLM for repair."""
        if llm:
            self.llm = llm
        elif OLLAMA_AVAILABLE:
            try:
                self.llm = Ollama(
                    base_url=settings.ollama_base_url,
                    model=settings.ollama_model,
                    temperature=0.1
                )
            except Exception as e:
                logger.warning(f"Failed to init Ollama: {e}")
                self.llm = None
        else:
            self.llm = None
        
        # Docling converter (lazy init)
        self._docling_converter = None
    
    @property
    def docling_converter(self) -> Optional[Any]:
        """Lazy initialization of Docling converter."""
        if not DOCLING_AVAILABLE:
            return None
        if self._docling_converter is None:
            self._docling_converter = DocumentConverter()
        return self._docling_converter
    
    def extract(self, file_path: str) -> ExtractionResult:
        """
        Main extraction method - orchestrates dual-path extraction.
        
        Args:
            file_path: Path to the file to extract
            
        Returns:
            ExtractionResult with text, images, and metadata
        """
        path = Path(file_path)
        ext = path.suffix.lower()
        
        # Check if image
        if ext in self.IMAGE_EXTENSIONS:
            return self._extract_image(file_path)
        
        # Check if document
        if ext not in self.DOCUMENT_EXTENSIONS:
            logger.warning(f"Unsupported file type: {ext}")
            return ExtractionResult(
                text=f"Unsupported file format: {ext}",
                source="error"
            )
        
        # Dual-Path Extraction for documents
        docling_result = self._extract_docling(file_path)
        native_result = self._extract_native(file_path)
        
        # Compare and merge
        return self._reconcile(docling_result, native_result)
    
    def _extract_docling(self, file_path: str) -> ExtractionResult:
        """Layer 1: Extract using Docling (structural/layout-aware)."""
        if not DOCLING_AVAILABLE or self.docling_converter is None:
            logger.warning("Docling not available, skipping Layer 1")
            return ExtractionResult(text="", source="docling_unavailable")
        
        try:
            result = self.docling_converter.convert(file_path)
            
            # Get markdown export (preserves structure)
            text = result.document.export_to_markdown()
            
            # Extract images if available
            images = []
            if hasattr(result.document, 'pictures'):
                for i, pic in enumerate(result.document.pictures):
                    images.append({
                        "index": i,
                        "caption": getattr(pic, 'caption', ''),
                        "data": getattr(pic, 'data', None)
                    })
            
            return ExtractionResult(
                text=text,
                images=images,
                metadata={"pages": getattr(result.document, 'num_pages', 0)},
                source="docling"
            )
            
        except Exception as e:
            logger.error(f"Docling extraction failed: {e}")
            return ExtractionResult(text="", source="docling_error")
    
    def _extract_native(self, file_path: str) -> ExtractionResult:
        """Layer 2: Extract using native specialized libraries."""
        path = Path(file_path)
        ext = path.suffix.lower()
        
        try:
            if ext == '.pdf':
                if PYPDF_AVAILABLE:
                    return self._extract_pdf_native(file_path)
                return ExtractionResult(text="Error: pypdf library not installed", source="missing_dependency_pdf")
            elif ext == '.docx':
                if DOCX_AVAILABLE:
                    return self._extract_docx_native(file_path)
                return ExtractionResult(text="Error: python-docx library not installed", source="missing_dependency_docx")
            elif ext == '.txt':
                return self._extract_txt(file_path)
            elif ext in ['.hwp', '.hwpx']:
                if HWP_AVAILABLE:
                    return self._extract_hwp_native(file_path)
                return ExtractionResult(text="Error: pyhwpx library not installed", source="missing_dependency_hwp")
            else:
                logger.warning(f"No native extractor for {ext}")
                return ExtractionResult(text="", source="native_unsupported")
                
        except Exception as e:
            logger.error(f"Native extraction failed: {e}")
            return ExtractionResult(text="", source="native_error")
    
    def _extract_pdf_native(self, file_path: str) -> ExtractionResult:
        """Extract PDF using pypdf."""
        reader = PdfReader(file_path)
        text_parts = []
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        return ExtractionResult(
            text="\n\n".join(text_parts),
            metadata={"pages": len(reader.pages)},
            source="native_pypdf"
        )
    
    def _extract_docx_native(self, file_path: str) -> ExtractionResult:
        """Extract DOCX using python-docx."""
        doc = DocxDocument(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        return ExtractionResult(
            text="\n".join(text_parts),
            metadata={"paragraphs": len(doc.paragraphs)},
            source="native_docx"
        )
    
    def _extract_hwp_native(self, file_path: str) -> ExtractionResult:
        """Extract HWP (한글) using pyhwpx or hwp_extract."""
        try:
            # Try pyhwpx first (more modern)
            try:
                from pyhwpx import Hwp
                hwp = Hwp()
                hwp.open(file_path)
                text = hwp.get_text()
                hwp.close()
                return ExtractionResult(
                    text=text,
                    metadata={"format": "hwp"},
                    source="native_pyhwpx"
                )
            except Exception:
                pass
            
            # Fallback to hwp_extract
            try:
                from hwp_extract import Hwp5File
                with Hwp5File(file_path) as hwp:
                    text_parts = []
                    for para in hwp.paragraphs():
                        if para.text.strip():
                            text_parts.append(para.text)
                    return ExtractionResult(
                        text="\n".join(text_parts),
                        metadata={"format": "hwp5"},
                        source="native_hwp_extract"
                    )
            except Exception:
                pass
            
            # Final fallback: try reading as binary and extracting text
            with open(file_path, 'rb') as f:
                content = f.read()
                # Simple text extraction from binary (basic fallback)
                text = content.decode('utf-8', errors='ignore')
                return ExtractionResult(
                    text=text,
                    metadata={"format": "hwp_binary"},
                    source="native_hwp_fallback"
                )
                
        except Exception as e:
            logger.error(f"HWP extraction failed: {e}")
            return ExtractionResult(text=f"HWP 추출 실패: {str(e)}", source="hwp_error")
    
    def _extract_txt(self, file_path: str) -> ExtractionResult:
        """Extract plain text file."""
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            text = f.read()
        
        return ExtractionResult(
            text=text,
            source="native_txt"
        )
    
    def _extract_image(self, file_path: str) -> ExtractionResult:
        """Extract text from image using OCR."""
        try:
            image = Image.open(file_path)
            # OCR extraction
            ocr_text = pytesseract.image_to_string(image, lang='kor+eng')
            
            return ExtractionResult(
                text=ocr_text,
                images=[{"path": file_path, "description": ""}],
                metadata={"size": image.size},
                source="ocr"
            )
        except Exception as e:
            logger.error(f"Image extraction failed: {e}")
            return ExtractionResult(text="", source="ocr_error")
    
    def _reconcile(
        self,
        docling_result: ExtractionResult,
        native_result: ExtractionResult
    ) -> ExtractionResult:
        """
        Compare and reconcile two extraction results.
        
        - High similarity: Use Docling (better formatting)
        - Low similarity: Use LLM to merge/repair
        """
        # If one is empty, use the other
        if not docling_result.text.strip():
            logger.info("Docling empty, using native result")
            return native_result
        
        if not native_result.text.strip():
            logger.info("Native empty, using docling result")
            return docling_result
        
        # Calculate similarity
        similarity = fuzz.ratio(
            docling_result.text[:5000],  # Limit for performance
            native_result.text[:5000]
        ) / 100.0
        
        logger.info(f"Extraction similarity: {similarity:.2%}")
        
        if similarity >= self.SIMILARITY_THRESHOLD:
            # High similarity: trust Docling's formatting
            logger.info("High similarity - using Docling result")
            result = ExtractionResult(
                text=docling_result.text,
                images=docling_result.images,
                metadata={
                    **docling_result.metadata,
                    "similarity": similarity,
                    "merge_method": "docling_preferred"
                },
                source="merged_high_similarity"
            )
            return result
        else:
            # Low similarity: use LLM to repair
            logger.info("Low similarity - triggering LLM repair")
            return self._llm_repair(docling_result, native_result, similarity)
    
    def _llm_repair(
        self,
        docling_result: ExtractionResult,
        native_result: ExtractionResult,
        similarity: float
    ) -> ExtractionResult:
        """Use LLM to reconcile divergent extractions."""
        
        if not self.llm:
            # Fallback: prefer docling if available, else native
            return docling_result if docling_result.text.strip() else native_result
        
        # Truncate for LLM context
        docling_text = docling_result.text[:3000]
        native_text = native_result.text[:3000]
        
        prompt = f"""다음은 동일한 문서에서 두 가지 방식으로 추출한 텍스트입니다.
유사도: {similarity:.1%}

[버전 A - 구조 추출 (레이아웃 인식)]
{docling_text}

[버전 B - 기본 추출 (Raw Text)]
{native_text}

두 버전을 비교하여 가장 정확한 최종 텍스트를 생성하세요.
- 숫자, 날짜, 이름 등 구체적인 정보는 버전 B를 우선
- 문단 구조와 형식은 버전 A를 참고
- 누락된 내용이 있으면 두 버전에서 보완

[최종 텍스트]"""

        try:
            repaired_text = self.llm.invoke(prompt)
            
            return ExtractionResult(
                text=repaired_text,
                images=docling_result.images,
                metadata={
                    "similarity": similarity,
                    "merge_method": "llm_repair",
                    "docling_len": len(docling_result.text),
                    "native_len": len(native_result.text)
                },
                source="merged_llm_repair"
            )
        except Exception as e:
            logger.error(f"LLM repair failed: {e}")
            # Fallback: concatenate both
            return ExtractionResult(
                text=f"[Docling]\n{docling_result.text}\n\n[Native]\n{native_result.text}",
                images=docling_result.images,
                metadata={"merge_method": "fallback_concat", "error": str(e)},
                source="merged_fallback"
            )


# Convenience function
def extract_document(file_path: str, llm: Optional[Ollama] = None) -> ExtractionResult:
    """
    Extract text and images from a document using dual-path extraction.
    
    Args:
        file_path: Path to the document
        llm: Optional LLM for repair (uses default if not provided)
        
    Returns:
        ExtractionResult with extracted content
    """
    extractor = DualPathExtractor(llm=llm)
    return extractor.extract(file_path)
